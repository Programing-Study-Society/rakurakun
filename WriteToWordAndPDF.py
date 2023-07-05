#Wordファイルを作成して、PDFファイルを作成する関数
#第一引数 タイトル部分のテキストが入ったリスト
#第二引数 本文部分のテキストが入ったリスト
#第三引数 画像ファイルの名前が入ったリスト
#第四引数 画像ファイルまでのパス(名前なし)
#第五引数 wordファイルの保存先
#第六引数 先生によるモード

import docx #ワードファイルを作成・編集・保存する為に使用
from docx.shared import Pt #ワードファイルのピクセルを変更する用
from docx.oxml.ns import qn #ワードファイルのフォントを変更する用
import pythoncom #COMランタイムの制御
import tkinter as tk #確認メッセージを表示する為に使用
import WriteToLog #ファイル分け
import WordToPDF #ファイル分け



def DocumentCreation():#Wordファイルの作成と初期設定をする関数
    #ドキュメントを作成
    pythoncom.CoInitialize()#COMランタイムを初期化
    doc = docx.Document()
    #ドキュメントのフォントを変更
    style = doc.styles['Normal']
    style.font.size = Pt(10.5)
    style.font.name = 'Times New Roman'
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '游明朝')
    style.font.name = '游明朝'
    
    return doc



def DocumentSaveAndPDFCreation(doc, savePath, log):#Wordファイルを保存して、エラーがなければPDF作成に移る関数
    try:
        doc.save(savePath + ".docx")  # Wordファイルを保存
        pythoncom.CoUninitialize()#COMランタイムの終了処理
    except PermissionError as e:
        WriteToLog.writeToLog(log,"ファイルの保存にアクセス権限がありません",'Red')
        WriteToLog.writeToLog(log,"対象ファイルが別プログラムから開かれている可能性があります",'Red')
        WriteToLog.writeToLog(log,"エラーメッセージ:" + str(e),'Red')
        return -1
    
    WordToPDF.word_to_pdf(savePath + ".docx", savePath+ ".pdf",log)#wordファイルをpdfファイルに変換して作成する


#ワードファイルを作成する関数
def wordWrite(text_titles, text_list, imageNames, imagePath, savePaths, mode, log):

    #阪口先生モードのWord作成
    if mode == "0":
        doc = DocumentCreation()

    #文字列のリストから、文字列を1つずつ取り出す
    for (text_title, text, savePath) in zip(text_titles, text_list, savePaths):

        #岩本先生モードのWord作成
        if mode == "1":
            doc = DocumentCreation()
        
        #タイトルを追加
        paragraph = doc.add_paragraph("課題" + text_title + "ソースコード")
        
        #改ページの処理（空白が出来ないように改行する、この一行のために一番苦労した）
        paragraph.paragraph_format.page_break_before=True
        
        #タイトルのフォントサイズを指定
        run = paragraph.runs[0]
        font = run.font
        font.size = Pt(20)  # ポイント単位で指定

        #本文を改行で分割してリストに挿入
        texts = text.split('\n')

        #本文を追加
        for text in texts:
           paragraph = doc.add_paragraph(text) #段落(パラグラフ)追加
           paragraph.paragraph_format.space_after = 0 #段落後のスペースを0にする
           paragraph.paragraph_format.line_spacing_rule = 0#行と段落の間隔を1.0Ptにする

        # paragraph = doc.add_paragraph(text) #本文を１つの段落でするならこっち

        #画像ファイルが1枚以上あれば画像挿入モードに移行する
        if(len(imageNames) > 0):

            #挿入する画像があるか検索する
            putFlag = False
            for img in imageNames:
                if((img.startswith(text_title)) and not(img[len(text_title)].isdigit())):
                    putFlag = True
                    break

            #今のソースコードに対応した画像ファイルがあれば続ける
            if putFlag:
                #実行結果という文を追加
                paragraph = doc.add_paragraph("課題" + text_title + "実行結果")

                #改ページの処理
                paragraph.paragraph_format.page_break_before=True

                #実行結果のフォントサイズを指定
                run = paragraph.runs[0]
                font = run.font
                font.size = Pt(20)  # ポイント単位で指定

                #画像を追加
                for img in imageNames:
                    if((img.startswith(text_title)) and not(img[len(text_title)].isdigit())) :# 挿入対象の画像なら
                        doc.add_picture(imagePath + img,width = Pt(300))#画像をWordに追

            #今のソースコードに対応した画像ファイルがなければ、そのことを知らせる
            else:
                #知らせをポップアップで出す
                tk.messagebox.showwarning("注意", text_title + "の画像ファイルが挿入されませんでした\n画像ファイルが見つからなかった可能性があります")

        #岩本先生モードのセーブ&PDF作成
        if mode == "1":
            errorChack = DocumentSaveAndPDFCreation(doc, savePath, log)
            if errorChack == -1 : return
    
    #阪口先生モードのセーブ&PDF作成
    if mode == "0":
        errorChack = DocumentSaveAndPDFCreation(doc, savePaths, log)
        if errorChack == -1 : return